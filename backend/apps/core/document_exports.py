from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from datetime import UTC, datetime
from html import escape
from io import BytesIO
from pathlib import PurePath
from uuid import UUID
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from django.db import transaction
from django.utils.text import slugify
from docx import Document as create_word_document
from docx.document import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from markdown_it import MarkdownIt
from markdown_it.token import Token

from .diagram_exports import (
    DiagramExportArtifact,
    diagram_manifest,
    embed_diagrams_in_html,
    render_diagram_exports,
)
from .document_attachments import copy_attachment_content
from .document_keys import KEY_TARGET_SCHEME, keys_in_markdown
from .documents import resolve_document
from .entity_mentions import resolve_entity_mentions
from .models import Document, DocumentAttachment
from .rendering import (
    RenderedAttachment,
    attachment_ids_in_markdown,
    entity_ids_in_markdown,
    render_markdown,
    render_pdf,
)
from .workspaces import ResolvedWorkspace

EXPORT_FORMATS = frozenset({"md", "html", "pdf", "docx", "bundle"})
BUNDLE_FORMAT = "tekdocs-portable-document/v1"
MAX_BUNDLE_ATTACHMENTS = 50
MAX_BUNDLE_ATTACHMENT_BYTES = 50 * 1024 * 1024
_MARKDOWN = MarkdownIt("commonmark", {"html": False}).enable(("table", "strikethrough"))
_FIXED_ZIP_TIME = (1980, 1, 1, 0, 0, 0)
_FIXED_DOCUMENT_TIME = datetime(2000, 1, 1, tzinfo=UTC)
_SAFE_EXTENSION = re.compile(r"^\.[A-Za-z0-9]{1,12}$")


class ExportConflict(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class ExportedAttachment:
    id: UUID
    filename: str
    media_type: str
    size: int
    checksum: str
    purpose: str
    version_number: int | None
    path: str
    content: bytes


@dataclass(frozen=True, slots=True)
class DocumentExportSnapshot:
    title: str
    markdown: str
    sanitized_html: str
    manifest: dict[str, object]
    digest: str
    attachments: tuple[ExportedAttachment, ...]
    diagrams: tuple[DiagramExportArtifact, ...]


def canonical_json(value: object) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode("utf-8")


def _attachment_path(record: DocumentAttachment) -> str:
    source = PurePath(record.original_filename)
    stem = slugify(source.stem)[:80] or "file"
    suffix = source.suffix.lower()
    extension = suffix if _SAFE_EXTENSION.fullmatch(suffix) else ""
    return f"attachments/{record.entity_id}/{stem}{extension}"


def _file_descriptor(path: str, media_type: str, content: bytes) -> dict[str, object]:
    return {
        "path": path,
        "media_type": media_type,
        "size": len(content),
        "checksum": hashlib.sha256(content).hexdigest(),
    }


@transaction.atomic
def resolve_export_snapshot(
    *, workspace: ResolvedWorkspace, document: Document, attachment_ids: tuple[UUID, ...] = ()
) -> DocumentExportSnapshot:
    """Freeze one authorized live composition and its requested managed files."""

    if len(attachment_ids) > MAX_BUNDLE_ATTACHMENTS:
        raise ExportConflict(f"A portable export may include at most {MAX_BUNDLE_ATTACHMENTS} files.")
    if len(set(attachment_ids)) != len(attachment_ids):
        raise ExportConflict("Each selected file may appear only once.")

    locked_document = (
        Document.objects.select_for_update(of=("self",))
        .select_related("entity", "organization", "organization__entity")
        .get(pk=document.pk)
    )
    resolved = resolve_document(locked_document)
    # Exports follow publication rather than the live view (ADR 0089): an exported
    # document has left the authorization boundary, so it may not carry a value that
    # was resolved for whoever happened to request the file. Until publish-time
    # resolution exists, a document with keys is refused rather than exported with
    # unresolved markers baked into bytes that outlive the request.
    keys, unparsable = keys_in_markdown(resolved.markdown)
    named = [key.expression for key in keys] + [
        target.removeprefix(KEY_TARGET_SCHEME) for target in unparsable
    ]
    if named:
        raise ExportConflict(
            "This document resolves keys from linked records, which cannot yet be exported: "
            f"{', '.join(sorted(set(named)))}."
        )
    requested_entities = entity_ids_in_markdown(resolved.markdown)
    entity_mentions = resolve_entity_mentions(workspace=workspace, markdown=resolved.markdown)
    if {UUID(entity_id) for entity_id in entity_mentions} != requested_entities:
        raise ExportConflict("The document contains an unavailable entity reference.")

    referenced_attachment_ids = attachment_ids_in_markdown(resolved.markdown)
    requested_files = set(attachment_ids)
    needed_attachment_ids = referenced_attachment_ids | requested_files
    attachment_records = list(
        DocumentAttachment.objects.select_for_update()
        .select_related("entity")
        .filter(
            document=locked_document,
            entity_id__in=needed_attachment_ids,
            archived_at__isnull=True,
            scan_status="clean",
        )
        .order_by("entity_id")
    )
    available_ids = {record.entity_id for record in attachment_records}
    if available_ids != needed_attachment_ids:
        raise ExportConflict("One or more selected or referenced files are unavailable.")

    rendered_attachments: dict[str, RenderedAttachment] = {}
    exported_attachments: list[ExportedAttachment] = []
    retained_bytes = 0
    for record in attachment_records:
        rendered_attachments[str(record.entity_id)] = {
            "id": str(record.entity_id),
            "filename": record.original_filename,
            "size": record.size,
        }
        if record.entity_id not in requested_files:
            continue
        try:
            content = copy_attachment_content(record)
        except Exception as exc:
            raise ExportConflict("A selected file failed its retained-content integrity check.") from exc
        retained_bytes += len(content)
        if retained_bytes > MAX_BUNDLE_ATTACHMENT_BYTES:
            raise ExportConflict("Portable export file content may not exceed 50 MiB.")
        exported_attachments.append(
            ExportedAttachment(
                id=record.entity_id,
                filename=record.original_filename,
                media_type=record.media_type,
                size=record.size,
                checksum=record.checksum,
                purpose=record.purpose,
                version_number=record.version_number,
                path=_attachment_path(record),
                content=content,
            )
        )

    diagrams = render_diagram_exports(resolved.markdown)
    sanitized_html = render_markdown(
        resolved.markdown,
        entity_mentions=entity_mentions,
        attachments=rendered_attachments,
    )
    sanitized_html = embed_diagrams_in_html(sanitized_html, diagrams)
    markdown_bytes = resolved.markdown.encode("utf-8")
    html_bytes = export_html(
        title=locked_document.entity.display_name,
        markdown=resolved.markdown,
        retained_html=sanitized_html,
    )
    attachment_manifest = [
        {
            "id": str(item.id),
            "filename": item.filename,
            "media_type": item.media_type,
            "size": item.size,
            "checksum": item.checksum,
            "purpose": item.purpose,
            "version_number": item.version_number,
            "referenced_in_markdown": item.id in referenced_attachment_ids,
            "path": item.path,
        }
        for item in exported_attachments
    ]
    files = [
        _file_descriptor("document/document.md", "text/markdown; charset=utf-8", markdown_bytes),
        _file_descriptor("document/document.html", "text/html; charset=utf-8", html_bytes),
        *[
            _file_descriptor(item.path, item.media_type or "application/octet-stream", item.content)
            for item in exported_attachments
        ],
        *[
            _file_descriptor(f"diagrams/{item.source.index:03d}.svg", "image/svg+xml", item.svg)
            for item in diagrams
            if item.svg is not None
        ],
        *[
            _file_descriptor(f"diagrams/{item.source.index:03d}.png", "image/png", item.png)
            for item in diagrams
            if item.png is not None
        ],
    ]
    organization = locked_document.organization
    manifest: dict[str, object] = {
        "format": BUNDLE_FORMAT,
        "export_class": "editable_revision_snapshot",
        "immutable_publication": False,
        "source_document_id": str(locked_document.entity_id),
        "workspace": {
            "kind": "organization" if organization is not None else "msp",
            "id": str(organization.entity_id) if organization is not None else None,
        },
        "title": locked_document.entity.display_name,
        "category": locked_document.category,
        "placements": [
            {
                "id": str(item.placement.id),
                "parent_id": str(item.placement.parent_id) if item.placement.parent_id else None,
                "position": item.placement.position,
                "depth": item.depth,
                "block_id": str(item.placement.block.entity_id),
                "resolution_mode": item.placement.resolution_mode,
                "revision_id": str(item.revision.id),
                "revision_number": item.revision.revision_number,
                "checksum": item.revision.checksum,
            }
            for item in resolved.placements
        ],
        "entities": [entity_mentions[entity_id] for entity_id in sorted(entity_mentions)],
        "attachments": attachment_manifest,
        "diagrams": diagram_manifest(diagrams),
        "files": files,
    }
    digest = hashlib.sha256(b"TEKDOCS-PORTABLE-EXPORT\x00v1\x00" + canonical_json(manifest)).hexdigest()
    manifest["content_digest"] = digest
    return DocumentExportSnapshot(
        title=locked_document.entity.display_name,
        markdown=resolved.markdown,
        sanitized_html=sanitized_html,
        manifest=manifest,
        digest=digest,
        attachments=tuple(exported_attachments),
        diagrams=diagrams,
    )


def export_html(
    *,
    title: str,
    markdown: str,
    retained_html: str | None = None,
    diagrams: tuple[DiagramExportArtifact, ...] = (),
) -> bytes:
    body = retained_html if retained_html is not None else render_markdown(markdown)
    if retained_html is None:
        body = embed_diagrams_in_html(body, diagrams)
    return (
        '<!doctype html>\n<html lang="en"><head><meta charset="utf-8">'
        '<meta http-equiv="Content-Security-Policy" '
        'content="default-src \'none\'; img-src data:; style-src \'unsafe-inline\'">'
        f"<title>{escape(title)}</title></head><body><main>{body}</main></body></html>\n"
    ).encode()


def export_pdf(*, title: str, markdown: str, diagrams: tuple[DiagramExportArtifact, ...] = ()) -> bytes:
    return render_pdf(markdown, title=title, diagrams=diagrams)


def _add_hyperlink(paragraph, text: str, url: str) -> None:  # type: ignore[no-untyped-def]
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "7A4A10")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((properties, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _write_inline(paragraph, children: list[Token]) -> None:  # type: ignore[no-untyped-def]
    bold = italic = strike = False
    index = 0
    while index < len(children):
        token = children[index]
        if token.type == "strong_open":
            bold = True
        elif token.type == "strong_close":
            bold = False
        elif token.type == "em_open":
            italic = True
        elif token.type == "em_close":
            italic = False
        elif token.type == "s_open":
            strike = True
        elif token.type == "s_close":
            strike = False
        elif token.type == "link_open":
            label: list[str] = []
            index += 1
            while index < len(children) and children[index].type != "link_close":
                if children[index].type in {"text", "code_inline"}:
                    label.append(children[index].content)
                index += 1
            url = str(token.attrGet("href") or "")
            if url.startswith(("http://", "https://", "mailto:")):
                _add_hyperlink(paragraph, "".join(label) or url, url)
            else:
                paragraph.add_run("".join(label) or url)
        elif token.type == "code_inline":
            run = paragraph.add_run(token.content)
            run.font.name = "Courier New"
        elif token.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
        elif token.type == "image":
            paragraph.add_run(token.content or token.attrGet("src") or "")
        elif token.type in {"text", "html_inline"}:
            run = paragraph.add_run(token.content)
            run.bold = bold
            run.italic = italic
            run.font.strike = strike
        index += 1


def _add_table(document: WordDocument, tokens: list[Token], start: int) -> int:
    rows: list[list[tuple[str, list[Token]]]] = []
    row: list[tuple[str, list[Token]]] | None = None
    cell_kind = "td"
    index = start + 1
    while index < len(tokens) and tokens[index].type != "table_close":
        token = tokens[index]
        if token.type == "tr_open":
            row = []
        elif token.type in {"th_open", "td_open"}:
            cell_kind = token.type[:2]
        elif token.type == "inline" and row is not None:
            row.append((cell_kind, token.children or []))
        elif token.type == "tr_close" and row is not None:
            rows.append(row)
            row = None
        index += 1
    if rows:
        column_count = max(len(row) for row in rows)
        table = document.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        for row_index, cells in enumerate(rows):
            for column_index, (kind, children) in enumerate(cells):
                paragraph = table.cell(row_index, column_index).paragraphs[0]
                _write_inline(paragraph, children)
                if kind == "th":
                    for run in paragraph.runs:
                        run.bold = True
    return index


def _normalize_docx_archive(content: bytes) -> bytes:
    source = ZipFile(BytesIO(content))
    target = BytesIO()
    with source, ZipFile(target, "w") as archive:
        for name in sorted(source.namelist()):
            info = ZipInfo(name, date_time=_FIXED_ZIP_TIME)
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            archive.writestr(info, source.read(name))
    return target.getvalue()


def export_docx(*, title: str, markdown: str, diagrams: tuple[DiagramExportArtifact, ...] = ()) -> bytes:
    document = create_word_document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.85)
    document.core_properties.title = title
    document.core_properties.author = "TekDocs"
    document.core_properties.last_modified_by = "TekDocs"
    document.core_properties.created = _FIXED_DOCUMENT_TIME
    document.core_properties.modified = _FIXED_DOCUMENT_TIME
    document.core_properties.revision = 1

    code_style = document.styles.add_style("TekDocs Code", WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    quote_style = document.styles.add_style("TekDocs Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote_style.font.italic = True
    document.add_heading(title, level=0)

    tokens = _MARKDOWN.parse(markdown)
    list_styles: list[str] = []
    heading_level: int | None = None
    blockquote_depth = 0
    diagram_index = 0
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_open":
            index = _add_table(document, tokens, index)
        elif token.type == "heading_open":
            heading_level = int(token.tag.removeprefix("h"))
        elif token.type == "heading_close":
            heading_level = None
        elif token.type == "bullet_list_open":
            list_styles.append("List Bullet")
        elif token.type == "ordered_list_open":
            list_styles.append("List Number")
        elif token.type in {"bullet_list_close", "ordered_list_close"}:
            if list_styles:
                list_styles.pop()
        elif token.type == "blockquote_open":
            blockquote_depth += 1
        elif token.type == "blockquote_close":
            blockquote_depth = max(0, blockquote_depth - 1)
        elif token.type == "inline":
            if heading_level is not None:
                paragraph = document.add_heading(level=heading_level)
            elif list_styles:
                style = list_styles[-1]
                paragraph = document.add_paragraph(style=style)
                if len(list_styles) > 1:
                    paragraph.paragraph_format.left_indent = Inches(0.25 * (len(list_styles) - 1))
            elif blockquote_depth:
                paragraph = document.add_paragraph(style="TekDocs Quote")
            else:
                paragraph = document.add_paragraph()
            _write_inline(paragraph, token.children or [])
        elif token.type == "fence" and token.info.strip().casefold() == "mermaid":
            item = diagrams[diagram_index] if diagram_index < len(diagrams) else None
            diagram_index += 1
            document.add_heading(item.source.title if item is not None else "Technical diagram", level=3)
            if item is not None and item.source.description:
                document.add_paragraph(item.source.description)
            if item is not None and item.png is not None:
                document.add_picture(BytesIO(item.png), width=Inches(6.2))
            else:
                document.add_paragraph("The diagram could not be rendered; its canonical source follows.")
            paragraph = document.add_paragraph(style="TekDocs Code")
            paragraph.add_run(token.content.rstrip("\n"))
        elif token.type in {"fence", "code_block"}:
            paragraph = document.add_paragraph(style="TekDocs Code")
            paragraph.add_run(token.content.rstrip("\n"))
        elif token.type == "hr":
            document.add_paragraph("―")
        index += 1

    output = BytesIO()
    document.save(output)
    return _normalize_docx_archive(output.getvalue())


def _zip_entry(name: str, content: bytes) -> tuple[ZipInfo, bytes]:
    info = ZipInfo(name, date_time=_FIXED_ZIP_TIME)
    info.compress_type = ZIP_DEFLATED
    info.external_attr = 0o600 << 16
    return info, content


def export_bundle(snapshot: DocumentExportSnapshot) -> bytes:
    entries = [
        ("manifest.json", canonical_json(snapshot.manifest) + b"\n"),
        (
            "document/document.html",
            export_html(
                title=snapshot.title,
                markdown=snapshot.markdown,
                retained_html=snapshot.sanitized_html,
            ),
        ),
        ("document/document.md", snapshot.markdown.encode("utf-8")),
        *((f"diagrams/{item.source.index:03d}.svg", item.svg) for item in snapshot.diagrams if item.svg is not None),
        *((f"diagrams/{item.source.index:03d}.png", item.png) for item in snapshot.diagrams if item.png is not None),
        *((item.path, item.content) for item in snapshot.attachments),
    ]
    output = BytesIO()
    with ZipFile(output, "w") as archive:
        for name, content in sorted(entries):
            info, encoded = _zip_entry(name, content)
            archive.writestr(info, encoded)
    return output.getvalue()
